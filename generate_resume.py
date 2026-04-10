from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

# Narrow margins
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

# --- NAME ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("JACK MORELLO")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
p.paragraph_format.space_after = Pt(2)

# --- CONTACT ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("(310) 902-5854  •  jaidanmorello@gmail.com  •  jackmorello.com  •  linkedin.com/in/jackmorello")
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p.paragraph_format.space_after = Pt(6)

# --- SUMMARY ---
p = doc.add_paragraph()
run = p.add_run("Relationship-driven strategist with 8+ years managing high-value client accounts, driving revenue growth, and building partnerships across agency, startup, and enterprise environments. Track record managing $2M+ in client budgets, advising Fortune 500 brands (Toyota, Disney, Amazon, Sony) and high-growth ventures through fundraising, product launches, and market expansion. Experienced translating complex business objectives into actionable strategy and delivering measurable outcomes for senior stakeholders.")
run.font.size = Pt(10)
p.paragraph_format.space_after = Pt(8)

def add_section_header(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): '999999'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_job(company, location, dates, title, description):
    p = doc.add_paragraph()
    run = p.add_run(company)
    run.bold = True
    run.font.size = Pt(10.5)
    run = p.add_run(f", {location}")
    run.font.size = Pt(10.5)
    # Add date on same line (tab right)
    run = p.add_run(f"\t{dates}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(6)

    p = doc.add_paragraph()
    run = p.add_run(title)
    run.italic = True
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(1)

    if description:
        p = doc.add_paragraph()
        run = p.add_run(description)
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.italic = True
        p.paragraph_format.space_after = Pt(3)

def add_bullet(bold_prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(f"{bold_prefix}: ")
        run.bold = True
        run.font.size = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)

# --- PROFESSIONAL EXPERIENCE ---
add_section_header("PROFESSIONAL EXPERIENCE")

add_job("INCUBELLA", "Los Angeles, CA", "Jun 2025 – Present", "Marketing Manager",
        "Growth studio partnering with early-stage ventures to build scalable business foundations.")
add_bullet("Client Advisory", "Serve as strategic advisor to founding teams, translating market dynamics into actionable growth plans aligned to each company's stage and objectives")
add_bullet("Partnership Development", "Source and onboard portfolio companies, designing tailored strategies across multiple concurrent client engagements")
add_bullet("Revenue Strategy", "Build acquisition and monetization frameworks, balancing resource allocation against portfolio-level outcomes")

add_job("NEAR", "Los Angeles, CA", "Apr 2022 – Jun 2024", "Marketing Manager",
        "AI-powered marketplace; led team through rapid scaling phase.")
add_bullet("Revenue Impact", "Led initiatives contributing to $1.8M+ in platform GMV, aligning strategy directly to business revenue targets")
add_bullet("Growth", "Scaled platform to 50K+ MAU through integrated multi-channel strategy")
add_bullet("Efficiency & ROI", "Reduced customer acquisition cost by 60% and improved conversion by 140% through structured testing frameworks")
add_bullet("Product Launches", "Managed go-to-market for 4 major releases, driving 15% adoption per launch by coordinating cross-functional stakeholders")

add_job("DTRAVEL", "Los Angeles, CA", "Apr 2021 – Apr 2022", "Head of Digital Growth",
        "Travel tech startup; led growth strategy through Series A fundraise and marketplace launch.")
add_bullet("Fundraising Support", "Built brand presence and community traction that directly supported the company's $8M Series A close")
add_bullet("Partner Acquisition", "Designed and executed partner onboarding funnel, acquiring 2,000+ property hosts to seed two-sided marketplace through consultative outreach at scale")
add_bullet("Community Building", "Grew engaged community from 7K to 55K+ organically, establishing trust and credibility in a new market")
add_bullet("Revenue Generation", "Conceived and launched digital collectible program generating $750K in incremental revenue")

add_job("DEVOLVED AI", "Woodland Hills, CA", "Sep 2024 – Jun 2025", "Marketing Manager",
        "AI platform; led positioning and market entry targeting enterprise buyers.")
add_bullet("Enterprise Positioning", "Developed messaging and go-to-market strategy targeting sophisticated enterprise buyers, translating complex technical capabilities into clear value propositions")
add_bullet("Demand Generation", "Built inbound pipeline generating qualified leads for sales team through thought leadership across LinkedIn and industry publications")
add_bullet("Brand & Conversion", "Led full rebrand and streamlined onboarding, reducing prospect drop-off by 35%")

add_job("IHEARTMEDIA (Unified)", "Los Angeles, CA", "Jan 2020 – Apr 2021", "Paid Media Manager",
        "Full-service agency managing enterprise accounts including Starzplay, Toyota, and 3+ additional brands.")
add_bullet("Client Relationship Management", "Day-to-day client lead for 5+ concurrent enterprise accounts, owning the relationship from strategy through executive-level QBR presentations")
add_bullet("Budget Stewardship", "Managed $200K–$400K/month in client budgets, making real-time allocation decisions to maximize performance against KPIs")
add_bullet("Performance Delivery", "Built optimization frameworks driving 25%+ ROAS improvement QoQ, consistently exceeding client benchmarks")
add_bullet("Stakeholder Communication", "Translated performance data into strategic insights for senior stakeholders, building trust and expanding account scope")

add_job("EPK.TV", "Los Angeles, CA", "Aug 2018 – Jan 2020", "Content Merchandiser",
        "Entertainment marketing platform serving major studios.")
add_bullet("Enterprise Client Operations", "Coordinated campaigns across 20+ studio partners including Disney, Amazon, Sony, NBC, and Netflix, managing high-volume deliverables across multiple stakeholder relationships")

add_job("24/7 LAUNDRY SERVICE (WASSERMAN)", "Los Angeles, CA", "May 2017 – Dec 2017", "Copywriter & Social Media Manager",
        "Social-first agency; clients included LG Mobile, Twitter, and Jordan Brand.")
add_bullet("Brand Partnerships", "Created content strategy and managed engagement for premium consumer brands, using performance data to refine approach and strengthen client outcomes")

# --- EDUCATION ---
add_section_header("EDUCATION")
p = doc.add_paragraph()
run = p.add_run("Bachelor of Science, Integrated Marketing Communications")
run.bold = True
run.font.size = Pt(10.5)
run = p.add_run("  |  Ithaca College, Ithaca, NY")
run.font.size = Pt(10.5)

# --- CORE COMPETENCIES ---
add_section_header("CORE COMPETENCIES")
p = doc.add_paragraph()
run = p.add_run("Client Strategy & Relationship Management  •  Partnership Development  •  Revenue Growth & P&L Ownership  •  Go-to-Market Strategy  •  Stakeholder Communication & Executive Presentations  •  Budget Management  •  Cross-Functional Leadership  •  Data-Driven Decision Making  •  Funnel Optimization")
run.font.size = Pt(10)
p.paragraph_format.space_after = Pt(4)

# --- TOOLS ---
add_section_header("TOOLS")
p = doc.add_paragraph()
run = p.add_run("Salesforce  •  HubSpot  •  Google Analytics (GA4)  •  Looker Studio  •  Tableau  •  Asana  •  Monday.com")
run.font.size = Pt(10)

# --- ADDITIONAL ---
add_section_header("ADDITIONAL")
p = doc.add_paragraph()
run = p.add_run("Co-Founder, Cribs.fun")
run.bold = True
run.font.size = Pt(10)
run = p.add_run(" — Mobile real estate app (product strategy, community-driven organic acquisition)")
run.font.size = Pt(10)

doc.save("/home/user/MORELLOSIMS/Jack_Morello_Resume_Partnerships.docx")
print("Done!")
